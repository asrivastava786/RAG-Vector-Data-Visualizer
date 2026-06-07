from dataclasses import dataclass, field
from html.parser import HTMLParser
from io import BytesIO


@dataclass
class ExtractedDocument:
    text: str
    structure: dict
    warnings: list[str] = field(default_factory=list)


class _HTMLTextParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        stripped = data.strip()
        if stripped:
            self.parts.append(stripped)


SUPPORTED_CONTENT_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
    "text/markdown",
    "text/html",
    "application/octet-stream",
}

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".html", ".htm"}


def extract_text(filename: str, content_type: str, payload: bytes) -> ExtractedDocument:
    suffix = _suffix(filename)
    if suffix == ".pdf" or content_type == "application/pdf":
        return _extract_pdf(payload)
    if suffix == ".docx" or content_type.endswith("wordprocessingml.document"):
        return _extract_docx(payload)
    if suffix in {".html", ".htm"} or content_type == "text/html":
        return _extract_html(payload)
    if suffix in {".txt", ".md"} or content_type.startswith("text/"):
        return _extract_plain_text(payload, suffix=suffix)
    return _extract_plain_text(payload, suffix=suffix, warning="Unknown type decoded as text.")


def validate_document_type(filename: str, content_type: str) -> None:
    from app.core.errors import bad_request

    suffix = _suffix(filename)
    if suffix not in SUPPORTED_EXTENSIONS and content_type not in SUPPORTED_CONTENT_TYPES:
        raise bad_request("Unsupported document type. Upload PDF, DOCX, TXT, MD, or HTML.")


def _suffix(filename: str) -> str:
    if "." not in filename:
        return ""
    return "." + filename.rsplit(".", 1)[-1].lower()


def _decode_text(payload: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "cp1252"):
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue
    return payload.decode("utf-8", errors="replace")


def _extract_plain_text(
    payload: bytes,
    *,
    suffix: str,
    warning: str | None = None,
) -> ExtractedDocument:
    text = _decode_text(payload)
    lines = text.splitlines()
    headings = [
        {"line": index + 1, "text": line.strip("# ").strip()}
        for index, line in enumerate(lines)
        if suffix == ".md" and line.lstrip().startswith("#")
    ]
    warnings = [warning] if warning else []
    return ExtractedDocument(
        text=text,
        structure={
            "kind": "markdown" if suffix == ".md" else "plain_text",
            "paragraph_count": len([line for line in lines if line.strip()]),
            "headings": headings,
            "tables": _detect_markdown_tables(lines),
        },
        warnings=warnings,
    )


def _extract_html(payload: bytes) -> ExtractedDocument:
    parser = _HTMLTextParser()
    parser.feed(_decode_text(payload))
    text = "\n\n".join(parser.parts)
    return ExtractedDocument(
        text=text,
        structure={"kind": "html", "paragraph_count": len(parser.parts), "headings": []},
    )


def _extract_pdf(payload: bytes) -> ExtractedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("PDF extraction dependency is not installed.") from exc

    reader = PdfReader(BytesIO(payload))
    page_texts = []
    for index, page in enumerate(reader.pages):
        page_texts.append({"page": index + 1, "text": page.extract_text() or ""})
    return ExtractedDocument(
        text="\n\n".join(page["text"] for page in page_texts),
        structure={
            "kind": "pdf",
            "pages": [
                {"page": page["page"], "character_count": len(page["text"])}
                for page in page_texts
            ],
            "page_count": len(page_texts),
            "headings": [],
        },
    )


def _extract_docx(payload: bytes) -> ExtractedDocument:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise RuntimeError("DOCX extraction dependency is not installed.") from exc

    document = DocxDocument(BytesIO(payload))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    headings = [
        {"index": index, "text": paragraph.text}
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.style and paragraph.style.name.lower().startswith("heading")
    ]
    tables = []
    for table_index, table in enumerate(document.tables):
        tables.append(
            {"index": table_index, "rows": len(table.rows), "columns": len(table.columns)}
        )
    return ExtractedDocument(
        text="\n\n".join(paragraphs),
        structure={
            "kind": "docx",
            "paragraph_count": len(paragraphs),
            "headings": headings,
            "tables": tables,
        },
    )


def _detect_markdown_tables(lines: list[str]) -> list[dict]:
    tables = []
    current_start: int | None = None
    for index, line in enumerate(lines):
        if "|" in line and line.strip().count("|") >= 2:
            current_start = index if current_start is None else current_start
        elif current_start is not None:
            tables.append({"start_line": current_start + 1, "end_line": index})
            current_start = None
    if current_start is not None:
        tables.append({"start_line": current_start + 1, "end_line": len(lines)})
    return tables
